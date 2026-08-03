"""
DOCX Layout Parser using python-docx.
Extracts headings, paragraphs, and tables from Word documents.
"""

import os
from typing import List
from docx import Document as DocxDocument
from loguru import logger

from backend.rag.parsers.base import BaseDocumentParser
from backend.rag.schemas import ElementType, ExtractedElement
from backend.utils.helpers import sanitize_text


class DOCXDocumentParser(BaseDocumentParser):
    """Layout-aware DOCX document parser."""

    def supports_extension(self, extension: str) -> bool:
        return extension.lower() in (".docx", ".doc")

    async def parse(self, file_path: str) -> List[ExtractedElement]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found at path: {file_path}")

        elements: List[ExtractedElement] = []
        heading_path: List[str] = []

        try:
            doc = DocxDocument(file_path)
            logger.info("Parsing DOCX document: path={path}", path=file_path)

            for para in doc.paragraphs:
                text = sanitize_text(para.text)
                if not text:
                    continue

                style_name = para.style.name.lower() if para.style else ""
                if "heading 1" in style_name:
                    element_type = ElementType.HEADING_1
                    heading_path = [text]
                elif "heading 2" in style_name:
                    element_type = ElementType.HEADING_2
                    heading_path = [heading_path[0], text] if heading_path else [text]
                elif "heading 3" in style_name:
                    element_type = ElementType.HEADING_3
                    heading_path.append(text)
                else:
                    element_type = ElementType.PARAGRAPH

                elements.append(
                    ExtractedElement(
                        element_type=element_type,
                        content=text,
                        page_number=1,
                        heading_path=list(heading_path),
                        metadata={"style_name": para.style.name if para.style else "Normal"},
                    )
                )

            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_grid = []
                for row in table.rows:
                    table_grid.append([cell.text.strip() for cell in row.cells])
                if table_grid:
                    table_md = self._format_table_as_markdown(table_grid)
                    elements.append(
                        ExtractedElement(
                            element_type=ElementType.TABLE,
                            content=table_md,
                            page_number=1,
                            heading_path=list(heading_path),
                            metadata={"table_index": table_idx},
                        )
                    )

            logger.info("DOCX parsing complete: path={path} elements={count}", path=file_path, count=len(elements))
            return elements
        except Exception as e:
            logger.error("Failed to parse DOCX file: path={path} error={err}", path=file_path, err=str(e))
            raise

    def _format_table_as_markdown(self, table_data: List[List[str]]) -> str:
        if not table_data:
            return ""
        lines = [f"| {' | '.join(table_data[0])} |", f"| {' | '.join('---' for _ in table_data[0])} |"]
        for row in table_data[1:]:
            lines.append(f"| {' | '.join(row)} |")
        return "\n".join(lines)
